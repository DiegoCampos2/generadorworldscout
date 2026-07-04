# -*- coding: utf-8 -*-

"""Word document (.docx) and PDF document generation module using python-docx and reportlab."""

import io
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Reportlab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage, Frame, PageTemplate
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# Hex colors based on brand
SCOUT_PURPLE = RGBColor(98, 37, 153)     # #622599
DARK_PURPLE = RGBColor(93, 47, 136)      # #5D2F88
GRAY_TEXT = RGBColor(100, 100, 100)
WHITE = RGBColor(255, 255, 255)

# Reportlab colors
RL_SCOUT_PURPLE = colors.HexColor("#622599")
RL_LIGHT_PURPLE = colors.HexColor("#2B0947")
RL_BORDER_COLOR = colors.HexColor("#D3D3D3")
RL_DATA_BORDER_COLOR = colors.HexColor("#E0E0E0")
RL_BG_ALT = colors.HexColor("#F8F9FA")

# ==============================================================================
# DOCX HELPER FUNCTIONS
# ==============================================================================

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) of a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, **kwargs):
    """Sets borders of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), edge_data.get('val', 'single'))
            border.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            border.set(qn('w:space'), str(edge_data.get('space', 0)))
            border.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(border)
        elif edge in kwargs:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
    tcPr.append(tcBorders)

def style_cell_header(cell, text, fill_hex="622599", font_size=10, text_color=WHITE, bold=True):
    """Applies standard header styling to a cell."""
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
    set_cell_borders(cell, 
        top={"sz": 4, "val": "single", "color": "D3D3D3"},
        bottom={"sz": 12, "val": "single", "color": "622599"},
        left={"sz": 4, "val": "single", "color": "D3D3D3"},
        right={"sz": 4, "val": "single", "color": "D3D3D3"}
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0] if p.runs else p.add_run(text)
    run.text = text
    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = text_color

def style_cell_data(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, font_size=10, text_color=None, fill_hex=None):
    """Applies standard data styling to a cell."""
    if fill_hex:
        set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    set_cell_borders(cell, 
        top={"sz": 4, "val": "single", "color": "E0E0E0"},
        bottom={"sz": 4, "val": "single", "color": "E0E0E0"},
        left={"sz": 4, "val": "single", "color": "E0E0E0"},
        right={"sz": 4, "val": "single", "color": "E0E0E0"}
    )
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    
    text_lines = text.split('\n')
    for idx, line in enumerate(text_lines):
        if idx > 0:
            p.add_run('\n')
        run = p.add_run(line)
        run.font.name = 'Arial'
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if text_color:
            run.font.color.rgb = text_color

# ==============================================================================
# DOCX GENERATOR STREAM
# ==============================================================================

ODS_IMAGES_DIR = os.path.join(os.path.dirname(__file__), 'ods_images')

LOGO_PATH = os.path.join(os.path.dirname(__file__), 'logos', 'wsc_logo_white.png')
TROOP_LOGO_PATH = os.path.join(os.path.dirname(__file__), 'logos', 'troop_compass.png')
BANNER_PATH = os.path.join(os.path.dirname(__file__), 'logos', 'banner.jpg')

def _add_docx_institutional_banner(doc):
    """Adds the custom banner image as header on every page."""
    if not os.path.exists(BANNER_PATH):
        return
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    # Clear default content
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run_img = p.add_run()
    run_img.add_picture(BANNER_PATH, width=Inches(6.5))


def generate_docx_stream(data):
    """
    Generates a Word Document in memory based on the structured data dictionary.
    Returns a io.BytesIO stream.
    """
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # Institutional Banner
    _add_docx_institutional_banner(doc)
    
    # Title Section
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    
    run_pre = title_p.add_run("Programa Scout: ")
    run_pre.font.name = 'Arial'
    run_pre.font.size = Pt(12)
    run_pre.font.bold = True
    run_pre.font.color.rgb = GRAY_TEXT
    
    run_main = title_p.add_run(f'“{data.get("title", "Actividad sin título")}”')
    run_main.font.name = 'Arial'
    run_main.font.size = Pt(16)
    run_main.font.bold = True
    run_main.font.color.rgb = SCOUT_PURPLE
    
    # Table 1: Fecha / Autor
    t1 = doc.add_table(rows=2, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_cell_header(t1.cell(0, 0), "Fecha")
    style_cell_header(t1.cell(0, 1), "Autor(es)")
    style_cell_data(t1.cell(1, 0), data.get("date", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell_data(t1.cell(1, 1), data.get("authors", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Table 2: Objetivo General / Especificos / Metas
    t2 = doc.add_table(rows=6, cols=1)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    style_cell_header(t2.cell(0, 0), "Objetivo General")
    style_cell_data(t2.cell(1, 0), data.get("general_objective", ""))
    
    style_cell_header(t2.cell(2, 0), "Objetivos específicos")
    spec_text = "\n".join([f"• {obj}" for obj in data.get("specific_objectives", []) if obj.strip()])
    if not spec_text:
        spec_text = "No se especificaron objetivos específicos."
    style_cell_data(t2.cell(3, 0), spec_text)
    
    style_cell_header(t2.cell(4, 0), "Metas")
    metas_text = "\n".join([f"• {meta}" for meta in data.get("goals", []) if meta.strip()])
    if not metas_text:
        metas_text = "No se especificaron metas."
    style_cell_data(t2.cell(5, 0), metas_text)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Table 3: Equipo Encargado / Adulto Acompañante
    t3 = doc.add_table(rows=2, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_cell_header(t3.cell(0, 0), "Equipo Encargado")
    style_cell_header(t3.cell(0, 1), "Adulto Acompañante")
    style_cell_data(t3.cell(1, 0), data.get("responsible_team", "Staff de la Comunidad"), align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell_data(t3.cell(1, 1), data.get("adult_companion", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Table 4: Lugar / Duracion / Cuota
    t4 = doc.add_table(rows=2, cols=3)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_cell_header(t4.cell(0, 0), "Lugar")
    style_cell_header(t4.cell(0, 1), "Duración")
    style_cell_header(t4.cell(0, 2), "Cuota")
    style_cell_data(t4.cell(1, 0), data.get("location", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell_data(t4.cell(1, 1), data.get("duration", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell_data(t4.cell(1, 2), data.get("fee", "0 Bs."), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Table 5: Indicadores de Logro / ODS
    t5 = doc.add_table(rows=3, cols=1)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    style_cell_header(t5.cell(0, 0), "Indicadores de Logro")
    logros_list = data.get("indicators", [])
    logros_text = " | ".join(logros_list) if logros_list else "No se seleccionaron indicadores."
    style_cell_data(t5.cell(1, 0), logros_text)
    
    style_cell_header(t5.cell(2, 0), "Objetivos de Desarrollo Sostenible")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    
    # ODS table with images
    ods_list = data.get("ods", [])
    if ods_list:
        ods_table = doc.add_table(rows=len(ods_list), cols=2)
        ods_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, item in enumerate(ods_list):
            ods_id = item.get("id", 0)
            name = item.get("name", "")
            justification = item.get("justification", "")
            # Image cell
            img_cell = ods_table.cell(idx, 0)
            img_cell.width = Inches(0.5)
            set_cell_margins(img_cell, top=40, bottom=40, left=60, right=60)
            set_cell_borders(img_cell,
                top={"sz": 4, "val": "single", "color": "E0E0E0"},
                bottom={"sz": 4, "val": "single", "color": "E0E0E0"},
                left={"sz": 4, "val": "single", "color": "E0E0E0"},
                right={"sz": 4, "val": "single", "color": "E0E0E0"}
            )
            p_img = img_cell.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ods_img_path = os.path.join(ODS_IMAGES_DIR, f"ods-{ods_id:02d}.png")
            if os.path.exists(ods_img_path):
                run_img = p_img.add_run()
                run_img.add_picture(ods_img_path, width=Inches(0.4))
            # Text cell
            text_cell = ods_table.cell(idx, 1)
            set_cell_margins(text_cell, top=40, bottom=40, left=80, right=80)
            set_cell_borders(text_cell,
                top={"sz": 4, "val": "single", "color": "E0E0E0"},
                bottom={"sz": 4, "val": "single", "color": "E0E0E0"},
                left={"sz": 4, "val": "single", "color": "E0E0E0"},
                right={"sz": 4, "val": "single", "color": "E0E0E0"}
            )
            p_text = text_cell.paragraphs[0]
            p_text.paragraph_format.line_spacing = 1.15
            run_name = p_text.add_run(f"{name}: ")
            run_name.font.name = 'Arial'
            run_name.font.size = Pt(10)
            run_name.font.bold = True
            run_text = p_text.add_run(justification)
            run_text.font.name = 'Arial'
            run_text.font.size = Pt(10)
    else:
        style_cell_data(t5.cell(3, 0) if t5.rows > 2 else t5.add_row().cells[0], "No se seleccionaron ODS.")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Header Cronograma
    cron_title = doc.add_paragraph()
    run_cron = cron_title.add_run("Cronograma de la actividad")
    run_cron.font.name = 'Arial'
    run_cron.font.size = Pt(12)
    run_cron.font.bold = True
    run_cron.font.color.rgb = SCOUT_PURPLE
    cron_title.paragraph_format.space_after = Pt(6)
    
    # Table 6: Cronograma
    schedule = data.get("schedule", [])
    t6 = doc.add_table(rows=1 + len(schedule), cols=4)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    style_cell_header(t6.cell(0, 0), "Hora")
    style_cell_header(t6.cell(0, 1), "Nombre de la actividad")
    style_cell_header(t6.cell(0, 2), "Materiales")
    style_cell_header(t6.cell(0, 3), "Encargado")
    
    for row_idx, item in enumerate(schedule):
        cell_idx = row_idx + 1
        style_cell_data(t6.cell(cell_idx, 0), item.get("time", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
        style_cell_data(t6.cell(cell_idx, 1), item.get("activity", ""))
        style_cell_data(t6.cell(cell_idx, 2), item.get("materials", "No aplica"))
        style_cell_data(t6.cell(cell_idx, 3), item.get("in_charge", ""))
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Header Dinámicas
    dyn_title = doc.add_paragraph()
    run_dyn = dyn_title.add_run("Descripciones de las dinámicas")
    run_dyn.font.name = 'Arial'
    run_dyn.font.size = Pt(12)
    run_dyn.font.bold = True
    run_dyn.font.color.rgb = SCOUT_PURPLE
    dyn_title.paragraph_format.space_after = Pt(6)
    
    # Table 7: Descripciones de las Dinámicas
    descriptions = data.get("descriptions", [])
    t7 = doc.add_table(rows=1 + len(descriptions), cols=2)
    t7.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_cell_header(t7.cell(0, 0), "Título de la Actividad")
    style_cell_header(t7.cell(0, 1), "Descripción")
    
    for row_idx, item in enumerate(descriptions):
        cell_idx = row_idx + 1
        style_cell_data(t7.cell(cell_idx, 0), item.get("title", ""), bold=True)
        style_cell_data(t7.cell(cell_idx, 1), item.get("description", ""))
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Header Escala DURAS
    scale_title = doc.add_paragraph()
    run_scale = scale_title.add_run("Escala de la reunión")
    run_scale.font.name = 'Arial'
    run_scale.font.size = Pt(12)
    run_scale.font.bold = True
    run_scale.font.color.rgb = SCOUT_PURPLE
    scale_title.paragraph_format.space_after = Pt(6)
    
    # Table 8: Escala DURAS
    t8 = doc.add_table(rows=6, cols=5)
    t8.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    style_cell_header(t8.cell(0, 0), "Escala")
    style_cell_header(t8.cell(0, 1), "Sí")
    style_cell_header(t8.cell(0, 2), "Regular")
    style_cell_header(t8.cell(0, 3), "No")
    style_cell_header(t8.cell(0, 4), "¿Por qué?")
    
    duras_keys = ["desafiante", "util", "recompensante", "atractiva", "segura"]
    duras_names = ["Desafiante", "Útil", "Recompensante", "Atractiva", "Segura"]
    duras_data = data.get("duras", {})
    
    for idx, key in enumerate(duras_keys):
        row_idx = idx + 1
        val_name = duras_names[idx]
        item = duras_data.get(key, {"value": "Regular", "reason": ""})
        selected_val = item.get("value", "Regular").strip().lower()
        
        style_cell_data(t8.cell(row_idx, 0), val_name, bold=True)
        
        val_si = "x" if selected_val in ["si", "sí", "yes"] else ""
        style_cell_data(t8.cell(row_idx, 1), val_si, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        val_reg = "x" if selected_val == "regular" else ""
        style_cell_data(t8.cell(row_idx, 2), val_reg, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        val_no = "x" if selected_val == "no" else ""
        style_cell_data(t8.cell(row_idx, 3), val_no, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        style_cell_data(t8.cell(row_idx, 4), item.get("reason", ""))
        
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

# ==============================================================================
# PDF GENERATOR STREAM (USING REPORTLAB)
# ==============================================================================

def _add_pdf_institutional_banner(story, cw, table_header_style, table_data_style):
    """Banner is now added via page template on every page - this function is a no-op."""
    pass


def generate_pdf_stream(data):
    """
    Generates a beautifully styled PDF Document in memory using ReportLab Platypus.
    Returns a io.BytesIO stream.
    """
    buffer = io.BytesIO()
    
    # Page dimensions setup (US Letter: 8.5 x 11 inches)
    margin = 0.8 * inch

    # Banner callback for every page
    def draw_banner(canvas, doc_obj):
        if os.path.exists(BANNER_PATH):
            canvas.saveState()
            # Draw banner image centered within page margins
            banner_w = letter[0] - 2 * margin
            banner_h = 1.2 * inch
            x = margin
            y = letter[1] - margin - banner_h
            canvas.drawImage(BANNER_PATH, x, y, width=banner_w, height=banner_h, preserveAspectRatio=True, anchor='n')
            canvas.restoreState()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin + 1.3*inch,
        bottomMargin=margin
    )
    
    styles = getSampleStyleSheet()
    
    # Modify/Define styles safely to avoid duplicate name crashes
    title_style = ParagraphStyle(
        'ScoutTitleStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=15,
        textColor=RL_SCOUT_PURPLE,
        alignment=TA_CENTER,
        spaceAfter=12
    )
    
    table_header_style = ParagraphStyle(
        'TableHeaderStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        textColor=colors.white,
        alignment=TA_CENTER
    )
    
    table_data_style = ParagraphStyle(
        'TableDataStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        textColor=colors.black,
        leading=11
    )
    
    table_data_center_style = ParagraphStyle(
        'TableDataCenterStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        textColor=colors.black,
        alignment=TA_CENTER,
        leading=11
    )
    
    section_header_style = ParagraphStyle(
        'SectionHeaderStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        textColor=RL_SCOUT_PURPLE,
        spaceBefore=10,
        spaceAfter=6
    )

    # Printable content width = 8.5 - 2*0.8 = 6.9 inches = 496.8 points
    cw = 496.8

    story = []
    
    # Institutional Banner
    _add_pdf_institutional_banner(story, cw, table_header_style, table_data_style)
    
    # 1. Title
    title_text = f"Programa Scout: \u201c{data.get('title', 'Actividad sin título')}\u201d"
    story.append(Paragraph(title_text, title_style))
    story.append(Spacer(1, 4))
    
    # 2. Table 1: Fecha / Autor
    t1_data = [
        [Paragraph("Fecha", table_header_style), Paragraph("Autor(es)", table_header_style)],
        [Paragraph(data.get("date", ""), table_data_center_style), Paragraph(data.get("authors", ""), table_data_center_style)]
    ]
    t1 = Table(t1_data, colWidths=[cw/2, cw/2])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), RL_SCOUT_PURPLE),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.5, RL_BORDER_COLOR),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t1)
    story.append(Spacer(1, 8))
    
    # 3. Table 2: Objetivos y Metas
    spec_text = "<br/>".join([f"• {obj}" for obj in data.get("specific_objectives", []) if obj.strip()])
    if not spec_text:
        spec_text = "No se especificaron objetivos específicos."
        
    metas_text = "<br/>".join([f"• {meta}" for meta in data.get("goals", []) if meta.strip()])
    if not metas_text:
        metas_text = "No se especificaron metas."
        
    t2_data = [
        [Paragraph("Objetivo General", table_header_style)],
        [Paragraph(data.get("general_objective", ""), table_data_style)],
        [Paragraph("Objetivos específicos", table_header_style)],
        [Paragraph(spec_text, table_data_style)],
        [Paragraph("Metas", table_header_style)],
        [Paragraph(metas_text, table_data_style)]
    ]
    t2 = Table(t2_data, colWidths=[cw])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,0), RL_SCOUT_PURPLE),
        ('BACKGROUND', (0,2), (0,2), RL_SCOUT_PURPLE),
        ('BACKGROUND', (0,4), (0,4), RL_SCOUT_PURPLE),
        ('GRID', (0,0), (-1,-1), 0.5, RL_BORDER_COLOR),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(t2)
    story.append(Spacer(1, 8))
    
    # 4. Table 3: Equipo / Adulto
    t3_data = [
        [Paragraph("Equipo Encargado", table_header_style), Paragraph("Adulto Acompañante", table_header_style)],
        [Paragraph(data.get("responsible_team", "Staff de la Comunidad"), table_data_center_style), Paragraph(data.get("adult_companion", ""), table_data_center_style)]
    ]
    t3 = Table(t3_data, colWidths=[cw/2, cw/2])
    t3.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), RL_SCOUT_PURPLE),
        ('GRID', (0,0), (-1,-1), 0.5, RL_BORDER_COLOR),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t3)
    story.append(Spacer(1, 8))
    
    # 5. Table 4: Lugar / Duración / Cuota
    t4_data = [
        [Paragraph("Lugar", table_header_style), Paragraph("Duración", table_header_style), Paragraph("Cuota", table_header_style)],
        [Paragraph(data.get("location", ""), table_data_center_style), Paragraph(data.get("duration", ""), table_data_center_style), Paragraph(data.get("fee", "0 Bs."), table_data_center_style)]
    ]
    t4 = Table(t4_data, colWidths=[cw/3, cw/3, cw/3])
    t4.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), RL_SCOUT_PURPLE),
        ('GRID', (0,0), (-1,-1), 0.5, RL_BORDER_COLOR),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t4)
    story.append(Spacer(1, 8))
    
    # 6. Table 5: Indicadores / ODS
    logros_list = data.get("indicators", [])
    logros_text = " | ".join(logros_list) if logros_list else "No se seleccionaron indicadores."
    
    ods_list = data.get("ods", [])
    ods_items = []
    for item in ods_list:
        name = item.get("name", "")
        justification = item.get("justification", "")
        ods_id = item.get("id", 0)
        if name:
            # Try to load ODS image
            ods_img_path = os.path.join(ODS_IMAGES_DIR, f"ods-{ods_id:02d}.png")
            if os.path.exists(ods_img_path):
                ods_img = RLImage(ods_img_path, width=0.35*inch, height=0.35*inch)
                ods_items.append([ods_img, Paragraph(f"<b>{name}:</b> {justification}", table_data_style)])
            else:
                ods_items.append([Paragraph(str(ods_id), table_data_center_style), Paragraph(f"<b>{name}:</b> {justification}", table_data_style)])
    
    # Build ODS table
    ods_header = [
        [Paragraph("ODS", table_header_style), Paragraph("Objetivos de Desarrollo Sostenible", table_header_style)]
    ]
    ods_rows = ods_header + ods_items if ods_items else [[Paragraph("No se seleccionaron ODS.", table_data_style)]]
    
    t_ods = Table(ods_rows, colWidths=[cw*0.08, cw*0.92])
    t_ods.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), RL_SCOUT_PURPLE),
        ('GRID', (0,0), (-1,-1), 0.5, RL_DATA_BORDER_COLOR),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(t_ods)
    
    # Page Break for better PDF structure
    story.append(PageBreak())
    
    # 7. Table 6: Cronograma
    story.append(Paragraph("Cronograma de la actividad", section_header_style))
    schedule = data.get("schedule", [])
    t6_data = [
        [Paragraph("Hora", table_header_style), Paragraph("Nombre de la actividad", table_header_style), Paragraph("Materiales", table_header_style), Paragraph("Encargado", table_header_style)]
    ]
    for item in schedule:
        t6_data.append([
            Paragraph(item.get("time", ""), table_data_center_style),
            Paragraph(item.get("activity", ""), table_data_style),
            Paragraph(item.get("materials", "No aplica"), table_data_style),
            Paragraph(item.get("in_charge", ""), table_data_center_style)
        ])
    
    # Column width distributions for 4-column layout
    t6 = Table(t6_data, colWidths=[cw*0.15, cw*0.45, cw*0.25, cw*0.15])
    t6_styles = [
        ('BACKGROUND', (0,0), (-1,0), RL_SCOUT_PURPLE),
        ('GRID', (0,0), (-1,-1), 0.5, RL_DATA_BORDER_COLOR),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]
    # Alternating row colors
    for r_idx in range(1, len(t6_data)):
        if r_idx % 2 == 0:
            t6_styles.append(('BACKGROUND', (0, r_idx), (-1, r_idx), RL_BG_ALT))
            
    t6.setStyle(TableStyle(t6_styles))
    story.append(t6)
    story.append(Spacer(1, 10))
    
    # 8. Table 7: Descripciones de Dinámicas
    descriptions = data.get("descriptions", [])
    if descriptions:
        story.append(Paragraph("Descripciones de las dinámicas", section_header_style))
        t7_data = [
            [Paragraph("Título de la Actividad", table_header_style), Paragraph("Descripción", table_header_style)]
        ]
        for item in descriptions:
            t7_data.append([
                Paragraph(f"<b>{item.get('title', '')}</b>", table_data_style),
                Paragraph(item.get("description", ""), table_data_style)
            ])
        t7 = Table(t7_data, colWidths=[cw*0.35, cw*0.65])
        t7_styles = [
            ('BACKGROUND', (0,0), (-1,0), RL_SCOUT_PURPLE),
            ('GRID', (0,0), (-1,-1), 0.5, RL_DATA_BORDER_COLOR),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]
        for r_idx in range(1, len(t7_data)):
            if r_idx % 2 == 0:
                t7_styles.append(('BACKGROUND', (0, r_idx), (-1, r_idx), RL_BG_ALT))
        t7.setStyle(TableStyle(t7_styles))
        story.append(t7)
        story.append(Spacer(1, 10))
        
    # 9. Table 8: Escala DURAS
    story.append(Paragraph("Escala de la reunión", section_header_style))
    t8_data = [
        [Paragraph("Escala", table_header_style), Paragraph("Sí", table_header_style), Paragraph("Regular", table_header_style), Paragraph("No", table_header_style), Paragraph("¿Por qué?", table_header_style)]
    ]
    
    duras_keys = ["desafiante", "util", "recompensante", "atractiva", "segura"]
    duras_names = ["Desafiante", "Útil", "Recompensante", "Atractiva", "Segura"]
    duras_data = data.get("duras", {})
    
    for idx, key in enumerate(duras_keys):
        val_name = duras_names[idx]
        item = duras_data.get(key, {"value": "Regular", "reason": ""})
        selected_val = item.get("value", "Regular").strip().lower()
        
        t8_data.append([
            Paragraph(f"<b>{val_name}</b>", table_data_style),
            Paragraph("X" if selected_val in ["si", "sí", "yes"] else "", table_data_center_style),
            Paragraph("X" if selected_val == "regular" else "", table_data_center_style),
            Paragraph("X" if selected_val == "no" else "", table_data_center_style),
            Paragraph(item.get("reason", ""), table_data_style)
        ])
        
    t8 = Table(t8_data, colWidths=[cw*0.22, cw*0.08, cw*0.12, cw*0.08, cw*0.50])
    t8_styles = [
        ('BACKGROUND', (0,0), (-1,0), RL_SCOUT_PURPLE),
        ('GRID', (0,0), (-1,-1), 0.5, RL_DATA_BORDER_COLOR),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]
    for r_idx in range(1, len(t8_data)):
        if r_idx % 2 == 0:
            t8_styles.append(('BACKGROUND', (0, r_idx), (-1, r_idx), RL_BG_ALT))
    t8.setStyle(TableStyle(t8_styles))
    story.append(t8)
    
    # Build Document PDF with banner on every page
    doc.build(story, onFirstPage=draw_banner, onLaterPages=draw_banner)
    
    buffer.seek(0)
    return buffer
